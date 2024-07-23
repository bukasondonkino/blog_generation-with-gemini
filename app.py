import io
import os
import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from PIL import Image
from dotenv import load_dotenv, find_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

load_dotenv(find_dotenv(), override=True)
os.environ.get('GOOGLE_API_KEY')
pexel_api_key = os.environ.get('PEXEL_API_KEY')

llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
    # other params...
    )


prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You are a helpful assistant that generate an article based on this {topic}"
        ),
        (
            "human",
            "{topic}"
        )
    ]
)

def fetch_photo(query):
    api_key = pexel_api_key
    url = "https://api.pexels.com/v1/search"
    
    headers = {
        "authorization": api_key
    }
    
    params = {
        'query': query,
        'per_page': 1
    }
    
    response = requests.get(url, headers = headers, params = params)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        else:
            st.write("No photos found for the given query.")
    else:
        st.write(f"Error: {response.status_code}, {response.text}")
        
    return None

def create_word_docx(user_input, paragraph, image_input):
    # Create a new Word document
    doc = Document()

    # Add the user input to the document
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)

    # Add the image to the document
    doc.add_heading('Image Input', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))  # Adjust the width as needed

    return doc

st.set_page_config(layout="wide")

def main():
    st.title("Article Generator App using Gemini 1.5 pro")

    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")

    image_input = st.text_input("Please enter the topic for the image you want to fetch!")
    
    if len(user_input) > 0 and len(image_input) > 0:
        col1, col2, col3 = st.columns([1,2,1])
        with col1:
            st.subheader("Generated content by Gemini")
            st.write("Topic of the article is: " + user_input)
            st.write("Image of the article is: " + image_input)
            chain = prompt | llm | StrOutputParser()
            result = chain.invoke({"topic": user_input})
            # print(result)
            if len(result) > 0:
                st.info("Your article has been been generated successfully!")
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")
                
        with col2:
            st.subheader("Fetched Image")
            image_url = fetch_photo(image_input)
            st.image(image_url)
            
        with col3:
            st.subheader("Final Article to Download")
            image_response = requests.get(image_url)
            img = Image.open(io.BytesIO(image_response.content))
            doc = create_word_docx(user_input, result, img)

            # Save the Word document to a BytesIO buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Prepare the download link
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )



if __name__ == "__main__":
    main()